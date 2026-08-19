"""
文档解析与向量化入库服务。

支持 PDF（PyPDF2）、Word（python-docx）、TXT、Markdown。
解析 → 分块 → 写入 Chroma 向量库，并更新处理进度。
"""

from __future__ import annotations

import hashlib
import io
import logging
from pathlib import Path
from typing import Callable

from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.models.knowledge import KnowledgeDocument
from app.services.vector_store import get_collection
from app.utils.text_chunker import chunk_text

logger = logging.getLogger(__name__)

ProgressCallback = Callable[[int, str], None]

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".xlsx"}


class DocumentProcessor:
    """知识库文档处理器。"""

    def __init__(self) -> None:
        self.upload_dir = Path(settings.KB_UPLOAD_DIR)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    @staticmethod
    def is_supported(filename: str) -> bool:
        return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS

    def parse_bytes(self, content: bytes, filename: str) -> str:
        """
        根据扩展名解析文档为纯文本。

        Raises:
            ValueError: 不支持的格式或解析失败
        """
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持 {SUPPORTED_EXTENSIONS}")

        try:
            if ext == ".pdf":
                return self._parse_pdf(content)
            if ext == ".docx":
                return self._parse_docx(content)
            if ext == ".xlsx":
                return self._parse_xlsx(content)
            if ext in (".txt", ".md"):
                return content.decode("utf-8", errors="ignore").strip()
        except Exception as exc:
            raise ValueError(f"文档解析失败 ({filename}): {exc}") from exc

        raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text.strip())
        text = "\n\n".join(p for p in pages if p)
        if not text:
            raise ValueError("PDF 未提取到文本（可能是扫描件）")
        return text

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        doc = DocxDocument(io.BytesIO(content))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        if not parts:
            raise ValueError("Word 文档内容为空")
        return "\n\n".join(parts)

    @staticmethod
    def _parse_xlsx(content: bytes) -> str:
        """将 xlsx 行为数据转为可检索文本（按景点聚合）。"""
        import pandas as pd

        df = pd.read_excel(io.BytesIO(content))
        if df.empty:
            raise ValueError("Excel 文件无数据")

        lines = [f"数据列: {', '.join(str(c) for c in df.columns)}"]
        name_col = "attraction_name" if "attraction_name" in df.columns else df.columns[4]

        # 优先索引 POI 列表中的灵山/拈花景点
        poi_path = Path(settings.SCENIC_POIS_FILE)
        if not poi_path.is_absolute():
            poi_path = Path(__file__).resolve().parents[2] / poi_path
        poi_names: set[str] = set()
        if poi_path.exists():
            import json

            data = json.loads(poi_path.read_text(encoding="utf-8"))
            poi_names = {p["name"] for p in data.get("pois", [])}

        if poi_names and name_col in df.columns:
            df = df[df[name_col].isin(poi_names)]

        if name_col not in df.columns:
            return lines[0]

        for name, group in df.groupby(name_col):
            block = [f"景点: {name}"]
            if "attraction_type" in group.columns:
                types = group["attraction_type"].dropna().unique()
                if len(types):
                    block.append(f"类型: {', '.join(str(t) for t in types[:3])}")
            if "attraction_content" in group.columns:
                sample = group["attraction_content"].dropna().astype(str).head(2).tolist()
                if sample:
                    block.append("内容摘要: " + "；".join(sample))
            if "stay_duration" in group.columns and group["stay_duration"].notna().any():
                block.append(f"平均停留(分钟): {group['stay_duration'].mean():.1f}")
            if "satisfaction" in group.columns and group["satisfaction"].notna().any():
                block.append(f"平均满意度: {group['satisfaction'].mean():.2f}")
            block.append(f"游客记录数: {len(group)}")
            lines.append("\n".join(block))

        text = "\n\n".join(lines)
        if len(text.strip()) < 20:
            raise ValueError("Excel 未提取到有效文本")
        return text

    def save_file(self, doc_id: str, filename: str, content: bytes) -> Path:
        """保存原始上传文件。"""
        safe_name = Path(filename).name
        path = self.upload_dir / f"{doc_id}_{safe_name}"
        path.write_bytes(content)
        return path

    def index_text_to_chroma(
        self,
        doc_id: str,
        title: str,
        filename: str,
        text: str,
        category: str,
        scenic_area_id: str,
    ) -> int:
        """
        文本分块并写入 Chroma（追加到已有集合，不 reset）。

        Returns:
            写入的 chunk 数量
        """
        collection = get_collection()

        # 删除该文档已有向量（支持重复上传覆盖）
        try:
            existing = collection.get(where={"doc_id": doc_id})
            if existing and existing.get("ids"):
                collection.delete(ids=existing["ids"])
        except Exception as exc:
            logger.warning("清理旧向量失败 doc_id=%s: %s", doc_id, exc)

        chunks = chunk_text(
            text,
            chunk_size=settings.RAG_CHUNK_SIZE,
            overlap=settings.RAG_CHUNK_OVERLAP,
        )
        if not chunks:
            raise ValueError("文档分块结果为空")

        ids: list[str] = []
        documents: list[str] = []
        metadatas: list[dict] = []

        for idx, chunk in enumerate(chunks):
            chunk_id = f"{doc_id}_{idx}"
            ids.append(chunk_id)
            documents.append(chunk)
            metadatas.append(
                {
                    "doc_id": doc_id,
                    "title": title,
                    "source_file": filename,
                    "category": category,
                    "scenic_area_id": scenic_area_id,
                    "chunk_index": idx,
                }
            )

        batch_size = 100
        for i in range(0, len(ids), batch_size):
            collection.add(
                ids=ids[i : i + batch_size],
                documents=documents[i : i + batch_size],
                metadatas=metadatas[i : i + batch_size],
            )

        return len(ids)

    async def process_document(
        self,
        db: AsyncSession,
        doc_id: str,
        filename: str,
        content: bytes,
        category: str,
        scenic_area_id: str,
    ) -> None:
        """
        完整处理流水线（后台任务调用），更新 DB 进度状态。

        状态流转：pending → parsing → chunking → indexing → ready / failed
        """
        title = Path(filename).stem

        async def update(progress: int, status: str, **extra: object) -> None:
            result = await db.execute(
                select(KnowledgeDocument).where(KnowledgeDocument.doc_id == doc_id)
            )
            doc = result.scalar_one_or_none()
            if doc:
                doc.progress = progress
                doc.status = status
                if "chunk_count" in extra:
                    doc.chunk_count = int(extra["chunk_count"])  # type: ignore[arg-type]
                if "error_message" in extra:
                    doc.error_message = extra["error_message"]  # type: ignore[assignment]
                await db.commit()
            logger.info("KB doc=%s status=%s progress=%d", doc_id, status, progress)

        try:
            await update(5, "parsing")
            text = self.parse_bytes(content, filename)

            await update(35, "chunking")
            self.save_file(doc_id, filename, content)

            await update(55, "indexing")
            chunk_count = self.index_text_to_chroma(
                doc_id=doc_id,
                title=title,
                filename=filename,
                text=text,
                category=category,
                scenic_area_id=scenic_area_id,
            )

            await update(100, "ready", chunk_count=chunk_count)
            logger.info("文档入库成功 doc_id=%s chunks=%d", doc_id, chunk_count)

        except Exception as exc:
            logger.exception("文档处理失败 doc_id=%s", doc_id)
            await update(0, "failed", error_message=str(exc))

    @staticmethod
    def make_doc_id(filename: str) -> str:
        raw = f"{filename}_{hashlib.md5(filename.encode()).hexdigest()[:8]}"
        return f"doc_{hashlib.md5(raw.encode()).hexdigest()[:12]}"
